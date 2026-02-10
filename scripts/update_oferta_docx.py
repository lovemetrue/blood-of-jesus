#!/usr/bin/env python3
"""
Скрипт для обновления договора оферты с новыми реквизитами самозанятого
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_oferta_docx(input_path, output_path):
    """Обновляет договор оферты с новыми реквизитами"""
    
    # Открываем документ
    doc = Document(input_path)
    
    # Данные для замены
    requisites_data = {
        "Получатель платежей": "Панов Дмитрий Александрович",
        "Статус": "Самозанятый",
        "ИНН самозанятого": "773273875610",
        "Банк-получатель": 'АО "ТБанк"',
        "БИК": "044525974",
        "Корр. счет": "30101810145250000974",
        "Расчетный счет": "40817810800000861767",
        "ИНН банка": "7710140679",
        "КПП": "771301001",
        "Валюта": "Российский рубль (RUB)",
        "Адрес": "г. Санкт-Петербург",
        "Телефон": "+7 994 417 89 86",
        "Email": "jesusthehealer@yandex.ru"
    }
    
    # Ищем и обновляем текст в документе
    for paragraph in doc.paragraphs:
        text = paragraph.text
        
        # Обновляем старые реквизиты на новые
        if "СБЕРБАНК" in text or "7707083893" in text or "40817810205001506701" in text:
            # Заменяем старые банковские данные
            text = text.replace("АСТРАХАНСКОЕ ОТДЕЛЕНИЕ N8625 ПАО СБЕРБАНК", 'АО "ТБанк"')
            text = text.replace("7707083893", "7710140679")
            text = text.replace("301502001", "771301001")
            text = text.replace("40817810205001506701", "40817810800000861767")
            text = text.replace("041203602", "044525974")
            text = text.replace("30101810500000000602", "30101810145250000974")
            text = text.replace("АСТРАХАНЬ", "Санкт-Петербург")
            text = text.replace("Астрахань", "Санкт-Петербург")
            
            # Удаляем назначение платежа, если есть
            if "Перевод средств по договору" in text or "5012452081" in text:
                # Удаляем строки с назначением платежа
                continue
            
            paragraph.text = text
    
    # Добавляем раздел с реквизитами в конец документа, если его там нет
    # Проверяем, есть ли уже раздел с реквизитами в конце
    last_text = " ".join([p.text for p in doc.paragraphs[-10:]])
    
    if "Реквизиты Исполнителя" not in last_text or "40817810800000861767" not in last_text:
        # Добавляем раздел с реквизитами
        doc.add_paragraph()
        doc.add_paragraph("8. Реквизиты Исполнителя").bold = True
        
        doc.add_paragraph()
        doc.add_paragraph("Получатель платежей").bold = True
        doc.add_paragraph(f"Получатель: {requisites_data['Получатель платежей']}")
        doc.add_paragraph(f"Статус: {requisites_data['Статус']}")
        doc.add_paragraph(f"ИНН самозанятого: {requisites_data['ИНН самозанятого']}")
        
        doc.add_paragraph()
        doc.add_paragraph("Банковские реквизиты").bold = True
        doc.add_paragraph(f"Банк-получатель: {requisites_data['Банк-получатель']}")
        doc.add_paragraph(f"БИК: {requisites_data['БИК']}")
        doc.add_paragraph(f"Корр. счет: {requisites_data['Корр. счет']}")
        doc.add_paragraph(f"Расчетный счет: {requisites_data['Расчетный счет']}")
        doc.add_paragraph(f"ИНН банка: {requisites_data['ИНН банка']}")
        doc.add_paragraph(f"КПП: {requisites_data['КПП']}")
        doc.add_paragraph(f"Валюта: {requisites_data['Валюта']}")
        
        doc.add_paragraph()
        doc.add_paragraph("Контактная информация").bold = True
        doc.add_paragraph(f"Адрес: {requisites_data['Адрес']}")
        doc.add_paragraph(f"Телефон: {requisites_data['Телефон']}")
        doc.add_paragraph(f"Email: {requisites_data['Email']}")
    
    # Сохраняем обновленный документ
    doc.save(output_path)
    print(f"✓ Документ обновлен: {output_path}")

if __name__ == "__main__":
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    
    input_file = project_root / "materials" / "oferta_773273875610.docx"
    output_file = project_root / "materials" / "oferta_773273875610.docx"
    
    if not input_file.exists():
        print(f"Ошибка: файл не найден: {input_file}")
        sys.exit(1)
    
    # Создаем резервную копию
    backup_file = project_root / "materials" / "oferta_773273875610_backup.docx"
    if not backup_file.exists():
        import shutil
        shutil.copy(input_file, backup_file)
        print(f"✓ Создана резервная копия: {backup_file}")
    
    update_oferta_docx(input_file, output_file)
    
    # Также обновляем файлы в публичных директориях
    public_files = [
        project_root / "public" / "oferta_773273875610.docx",
        project_root / "frontend" / "public" / "oferta_773273875610.docx"
    ]
    
    for pub_file in public_files:
        if pub_file.exists():
            import shutil
            shutil.copy(output_file, pub_file)
            print(f"✓ Обновлен файл в публичной директории: {pub_file}")
